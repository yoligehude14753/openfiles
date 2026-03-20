from pathlib import Path
from typing import Dict, Any
import docx
from .base import BaseParser


class DocxParser(BaseParser):
    SUPPORTED_TYPES = ["docx"]

    def can_parse(self, file_type: str) -> bool:
        return file_type in self.SUPPORTED_TYPES

    def parse(self, file_path: Path) -> Dict[str, Any]:
        try:
            doc = docx.Document(file_path)

            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

            tables_text = []
            for table in doc.tables:
                rows = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    rows.append(" | ".join(cells))
                if rows:
                    tables_text.append("\n".join(rows))

            parts = []
            if paragraphs:
                parts.append("\n\n".join(paragraphs))
            if tables_text:
                parts.append("\n\n[Tables]\n" + "\n\n".join(tables_text))

            full_text = "\n\n".join(parts)

            return {
                "content": full_text,
                "paragraphs": paragraphs,
                "num_paragraphs": len(paragraphs),
                "num_tables": len(tables_text),
                "success": True,
            }
        except Exception as e:
            return {"content": "", "error": str(e), "success": False}
